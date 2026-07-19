import streamlit as st
from docxtpl import DocxTemplate, RichText
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import os
import re 
import json
from num2words import num2words
import pandas as pd
from datetime import datetime
import gspread
from google.oauth2.service_account import Credentials
import requests 

import docx
import copy

try:
    from google import genai
    from PIL import Image
    HAS_AI = True
except ImportError:
    HAS_AI = False

TEMPLATE_NAME = "образец отчета.docx"

st.set_page_config(page_title="Генератор Отчетов - Гарант Оценка", layout="wide")

if HAS_AI and "GEMINI_API_KEY" in st.secrets:
    try:
        gemini_client = genai.Client(api_key=st.secrets["GEMINI_API_KEY"])
        AI_READY = True
    except Exception as e:
        AI_READY = False
        st.sidebar.error(f"⚠️ ИИ-сканер недоступен: {e}")
else:
    AI_READY = False

KG_REGIONS = {
    "г. Бишкек": ["Ленинский район", "Октябрьский район", "Первомайский район", "Свердловский район"],
    "г. Ош": ["Центральный", "Амир-Тимур", "Толойкон", "Керме-Тоо", "Жапалак"],
    "Чуйская область": ["Аламудунский район", "Ысык-Атинский район", "Сокулукский район", "Московский район", "Панфиловский район", "Жайылский район", "Кеминский район", "Чуйский район", "г. Токмок"],
    "Ошская область": ["Кара-Сууский район", "Ноокатский район", "Узгенский район", "Алайский район", "Араванский район", "Чон-Алайский район", "Кара-Кулджинский район"],
    "Джалал-Абадская область": ["Сузакский район", "Базар-Коргонский район", "Ноокенский район", "Аксыйский район", "Ала-Букинский район", "Чаткальский район", "Токтогульский район", "Тогуз-Тороуский район", "г. Джалал-Абад", "г. Кара-Куль", "г. Таш-Кумыр", "г. Майлуу-Суу"],
    "Иссык-Кульская область": ["Иссык-Кульский район", "Тюпский район", "Ак-Суйский район", "Джети-Огузский район", "Тонский район", "г. Каракол", "г. Балыкчы"],
    "Нарынская область": ["Нарынский район", "Ат-Башинский район", "Ак-Талинский район", "Жумгальский район", "Кочкорский район", "г. Нарын"],
    "Баткенская область": ["Баткенский район", "Кадамжайский район", "Лейлекский район", "г. Баткен", "г. Кызыл-Кыя", "г. Сулюкта"],
    "Таласская область": ["Таласский район", "Бакай-Атинский район", "Кара-Бууринский район", "Манасский район", "г. Талас"]
}

def upload_to_telegram(file_bytes, file_name):
    try:
        token = st.secrets.get("TELEGRAM_BOT_TOKEN", "")
        chat_id = st.secrets.get("TELEGRAM_CHAT_ID", "")
        
        if not token or not chat_id:
            st.error("❌ Ошибка: В Secrets не добавлены настройки Telegram!")
            return "Ошибка настроек"
            
        url = f"https://api.telegram.org/bot{token}/sendDocument"
        files = {"document": (file_name, io.BytesIO(file_bytes), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        data = {"chat_id": chat_id, "caption": f"📄 Сгенерирован новый отчет!\n🚗 Автомобиль: {file_name.replace('.docx', '')}\n📅 Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}"}
        
        response = requests.post(url, data=data, files=files)
        if response.json().get("ok"):
            st.info("🚀 Отчет успешно отправлен в Telegram-чат!")
            return "Отправлено в Telegram"
        else:
            return "Ошибка отправки"
    except Exception as e:
        return "Сбой системы"

def get_google_sheets_client():
    scopes = ["https://www.googleapis.com/auth/spreadsheets", "https://www.googleapis.com/auth/drive"]
    creds = Credentials.from_service_account_info(st.secrets["gcp_service_account"], scopes=scopes)
    return gspread.authorize(creds)

def append_to_google_sheets(boss_row, db_row):
    try:
        client = get_google_sheets_client()
        doc = client.open_by_key(st.secrets["spreadsheet_id"])
        doc.get_worksheet(0).append_row(boss_row)
        try:
            sheet_db = doc.worksheet("База_проверок")
        except gspread.exceptions.WorksheetNotFound:
            sheet_db = doc.add_worksheet(title="База_проверок", rows="1000", cols="10")
            sheet_db.append_row(["Номер отчета", "Госномер", "VIN код", "Техпаспорт", "Дата отчета", "Статус файла"])
        sheet_db.append_row(db_row)
        return True
    except:
        return False

def _sanitize_for_display(df):
    """Приводит все колонки к единому типу, чтобы pandas/pyarrow не падали
    при смешанных типах данных (текст+числа) в одной колонке Google Sheets."""
    if df is None or df.empty:
        return df
    df = df.copy()
    for col in df.columns:
        if df[col].dtype == object:
            df[col] = df[col].apply(lambda x: '' if x is None else str(x))
    return df

@st.cache_data(ttl=60)
def get_cached_preview():
    try:
        raw = pd.DataFrame(get_google_sheets_client().open_by_key(st.secrets["spreadsheet_id"]).get_worksheet(0).get_all_records())
        return _sanitize_for_display(raw)
    except: return None

@st.cache_data(ttl=60)
def get_cached_db():
    try:
        raw = pd.DataFrame(get_google_sheets_client().open_by_key(st.secrets["spreadsheet_id"]).worksheet("База_проверок").get_all_records())
        return _sanitize_for_display(raw)
    except: return pd.DataFrame()

def force_font_everywhere(document, font_name='Times New Roman'):
    """Принудительно ставит один шрифт на ВЕСЬ документ целиком: основной текст
    шаблона, все таблицы (включая вложенные в фотоотчет и смету, которые
    подмешиваются docxtpl как subdoc-и), а также колонтитулы."""
    def set_run_font(r):
        r.font.name = font_name
        rPr = r._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
            rFonts.set(qn(attr), font_name)

    def process_paragraphs(paragraphs):
        for p in paragraphs:
            for r in p.runs:
                set_run_font(r)

    def process_tables(tables):
        for t in tables:
            for row in t.rows:
                for cell in row.cells:
                    process_paragraphs(cell.paragraphs)
                    process_tables(cell.tables)

    process_paragraphs(document.paragraphs)
    process_tables(document.tables)
    for section in document.sections:
        for hf in (section.header, section.footer, section.first_page_header,
                   section.first_page_footer, section.even_page_header, section.even_page_footer):
            process_paragraphs(hf.paragraphs)
            process_tables(hf.tables)

def parse_sum_from_text(text):
    clean_text = text.replace('\xa0', ' ')
    match = re.search(r'(\d[\d\s]*[.,]\d+)', clean_text)
    if match:
        val_str = match.group(1).replace(' ', '').replace(',', '.')
        try: return float(val_str)
        except: return 0.0
    return 0.0

def format_sum(val):
    return f"{val:,.2f}".replace(',', 'X').replace('.', ',').replace('X', ' ')

def find_duplicate_matches(dfs, checks):
    """Ищет точные совпадения непустых значений (техпаспорт, гос.номер, VIN,
    номер отчета, ФИО заказчика) в любой из колонок ранее сохраненных отчетов,
    чтобы предупредить об уже существующем/повторном отчете."""
    warnings = []
    for field_label, value in checks:
        value_norm = (value or "").strip().lower()
        if not value_norm:
            continue
        for sheet_label, df in dfs:
            if df is None or df.empty:
                continue
            for col in df.columns:
                try:
                    matches = df[df[col].astype(str).str.strip().str.lower() == value_norm]
                except Exception:
                    continue
                if not matches.empty:
                    for _, row in matches.iterrows():
                        row_preview = ", ".join(f"{k}: {v}" for k, v in row.items() if str(v).strip())
                        warnings.append(f"⚠️ **{field_label}** «{value}» уже встречается в «{sheet_label}» (колонка «{col}») — {row_preview}")
                    break  # одной найденной колонки в этом листе достаточно
    return warnings

def process_smart_calc_tables(uploaded_file, ext_serv="", ext_parts="", ext_mat=""):
    try:
        doc_in = docx.Document(uploaded_file)
        doc_out = docx.Document()
        tables_found = {}
        
        def format_table_smart(tbl_element):
            tblPr = tbl_element.find(qn('w:tblPr'))
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl_element.insert(0, tblPr)
            tblW = tblPr.find(qn('w:tblW'))
            if tblW is None:
                tblW = OxmlElement('w:tblW')
                tblPr.append(tblW)
            tblW.set(qn('w:w'), '5000')
            tblW.set(qn('w:type'), 'pct')
            
            rows = tbl_element.findall(qn('w:tr'))
            if len(rows) > 0:
                for row_idx, tr in enumerate(rows):
                    for tc in tr.findall(qn('w:tc')):
                        for p in tc.findall(qn('w:p')):
                            for r in p.findall(qn('w:r')):
                                rPr = r.find(qn('w:rPr'))
                                if rPr is None:
                                    rPr = OxmlElement('w:rPr')
                                    r.insert(0, rPr)

                                # Шрифт Times New Roman для ВСЕХ строк таблицы (не только заголовка)
                                rFonts = rPr.find(qn('w:rFonts'))
                                if rFonts is None:
                                    rFonts = OxmlElement('w:rFonts')
                                    rPr.append(rFonts)
                                for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
                                    rFonts.set(qn(attr), 'Times New Roman')

                                # Жирный шрифт только для строки заголовка (первая строка)
                                if row_idx == 0:
                                    b = rPr.find(qn('w:b'))
                                    if b is None:
                                        b = OxmlElement('w:b')
                                        rPr.append(b)

            return tbl_element

        def apply_font(run):
            run.font.name = 'Times New Roman'
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
                rFonts.set(qn(attr), 'Times New Roman')

        for table in doc_in.tables:
            prev_elm = table._element.getprevious()
            header_text = ""
            while prev_elm is not None:
                if prev_elm.tag.endswith('p'):
                    p = docx.text.paragraph.Paragraph(prev_elm, doc_in)
                    if p.text.strip():
                        header_text = p.text.lower()
                        break
                prev_elm = prev_elm.getprevious()
            
            col_headers = [cell.text.lower() for cell in table.rows[0].cells] if table.rows else []
            
            if "воздейств" in header_text or "затрат" in header_text or "услуг" in header_text or "нормо-час" in col_headers:
                tables_found['services'] = table
            elif "запасных" in header_text or "запчаст" in header_text:
                tables_found['parts'] = table
            elif "материал" in header_text:
                tables_found['materials'] = table

        has_any = False
        approval_lines = []
        overall_total = 0.0
        
        if 'services' in tables_found:
            has_any = True
            p_title = doc_out.add_paragraph()
            r_title = p_title.add_run("Перечень и стоимость затрат (услуг), необходимых для восстановления:")
            r_title.bold = True
            apply_font(r_title)
            
            copied_tbl = copy.deepcopy(tables_found['services']._element)
            copied_tbl = format_table_smart(copied_tbl)
            p_title._p.addnext(copied_tbl)
            
            p_note = doc_out.add_paragraph()
            r_note = p_note.add_run("Примечание: ")
            r_note.bold = True
            apply_font(r_note)
            note_text = "стоимость нормо-часа ремонтно-восстановительных работ (1500,00 сом) определена согласно анализу стоимости услуг на станциях технического обслуживания: ИП «Сергей», +996702200885; +996553535533; +996550444488, +996559885102, +996550180555; +996555495545."
            if ext_serv.strip():
                note_text += f"\nДополнительно: {ext_serv.strip()}"
            r_note_body = p_note.add_run(note_text + "\n")
            apply_font(r_note_body)
            
            last_row_text = " ".join([cell.text for cell in tables_found['services'].rows[-1].cells])
            val = parse_sum_from_text(last_row_text)
            if val > 0:
                approval_lines.append(f"Услуг – {format_sum(val)} сом;")
                overall_total += val

        if 'parts' in tables_found:
            has_any = True
            p_title = doc_out.add_paragraph()
            r_title = p_title.add_run("Стоимость запасных частей:")
            r_title.bold = True
            apply_font(r_title)
            
            copied_tbl = copy.deepcopy(tables_found['parts']._element)
            copied_tbl = format_table_smart(copied_tbl)
            p_title._p.addnext(copied_tbl)
            
            p_note = doc_out.add_paragraph()
            r_note = p_note.add_run("Примечание: ")
            r_note.bold = True
            apply_font(r_note)
            note_text = "указана средняя стоимость запасных частей, поддержанных оригинальных, дубликатов на основании анализа рынка ЕАЭС.\nСсылки: в качестве информации была использована база данных ОсОО «Гарант Оценка»; интернет-ресурсы: mashina.kg, lalafo.kg; +996 551 411 711; +996 504 386 999; +996 500 524 624; +996 556 522 516; +996 707 008 833; +996 707 380 001."
            if ext_parts.strip():
                note_text += f"\nДополнительные ссылки: {ext_parts.strip()}"
            r_note_body = p_note.add_run(note_text + "\n")
            apply_font(r_note_body)
            
            last_row_text = " ".join([cell.text for cell in tables_found['parts'].rows[-1].cells])
            val = parse_sum_from_text(last_row_text)
            if val > 0:
                approval_lines.append(f"Запасных частей – {format_sum(val)} сом;")
                overall_total += val

        if 'materials' in tables_found:
            has_any = True
            p_title = doc_out.add_paragraph()
            r_title = p_title.add_run("Стоимость материалов:")
            r_title.bold = True
            apply_font(r_title)
            
            copied_tbl = copy.deepcopy(tables_found['materials']._element)
            copied_tbl = format_table_smart(copied_tbl)
            p_title._p.addnext(copied_tbl)
            
            p_note = doc_out.add_paragraph()
            r_note = p_note.add_run("Примечание: ")
            r_note.bold = True
            apply_font(r_note)
            note_text = "указана средняя стоимость материалов, источники конъюнктурного анализа: +996 708 707 332; +996 13 54 46; +996 550 98 77 01; +996 553 40 03 98"
            if ext_mat.strip():
                note_text += f"\nДополнительно: {ext_mat.strip()}"
            r_note_body = p_note.add_run(note_text + "\n")
            apply_font(r_note_body)
            
            last_row_text = " ".join([cell.text for cell in tables_found['materials'].rows[-1].cells])
            val = parse_sum_from_text(last_row_text)
            if val > 0:
                approval_lines.append(f"Материалов – {format_sum(val)} сом;")
                overall_total += val
            
        if not has_any:
            doc_out.add_paragraph("Таблицы с расчетами не найдены в загруженном документе.")
            
        buffer = io.BytesIO()
        doc_out.save(buffer)
        buffer.seek(0)
        
        approval_text = "\n".join(approval_lines)
        if overall_total > 0:
            approval_text += f"\nИтого: {format_sum(overall_total)} сом."
            
        return buffer, approval_text
    except Exception as e:
        st.error(f"❌ Ошибка обработки таблиц сметы: {e}")
        return None, ""

DEFAULT_DAMAGE_SUFFIX = "Дефектный акт на транспортное средство на дату оценки не предоставлялся. Оценка технического состояния произведена без учёта скрытых дефектов."
DEFAULT_REPAIR_SUFFIX = "После завершения ремонтно-восстановительных работ необходим контроль геометрии кузова, зазоров навесных элементов и качества ЛКП. Контроль выполняется организацией, осуществляющей ремонт."

def clear_fields():
    fields_to_clear = [
        "report_num", "contract_num", "date_ocenki", "customer",
        "aymak_input", "street_address", "sum_num", "car_model", 
        "reg_num", "vin", "tech_passport", "year", "engine_vol", 
        "color", "body_type", "service_cost", "extra_services", "extra_parts", "extra_materials"
    ]
    for f in fields_to_clear:
        if f in st.session_state:
            st.session_state[f] = ""
    
    if "steering" in st.session_state:
        st.session_state["steering"] = "Левый руль"
    if "date_otcheta" in st.session_state:
        st.session_state["date_otcheta"] = datetime.now().strftime("%d.%m.%Y")
    if "region_select" in st.session_state:
        st.session_state["region_select"] = list(KG_REGIONS.keys())[0]
    if "district_select" in st.session_state:
        st.session_state["district_select"] = KG_REGIONS[list(KG_REGIONS.keys())[0]][0]
        
    st.session_state.damage_text = DEFAULT_DAMAGE_SUFFIX
    st.session_state.repair_text = f"Для восстановления требуется выполнить комплекс слесарно-кузовных, рихтовочных и малярно-окрасочных работ с применением расходных материалов, с последующей сборкой и регулировкой навесных элементов.\n{DEFAULT_REPAIR_SUFFIX}"

st.title("🚗 Главное рабочее место оценщика")

if os.path.exists(TEMPLATE_NAME):
    st.success(f"✅ Базовый шаблон отчета (`{TEMPLATE_NAME}`) подключен.")
    template_source = TEMPLATE_NAME
else:
    template_source = st.file_uploader("Загрузите шаблон отчета", type="docx")

if AI_READY:
    with st.expander("🤖 Умный сканер техпаспорта", expanded=True):
        sts_images = st.file_uploader("Загрузить фото техпаспорта", type=["jpg", "jpeg", "png"], accept_multiple_files=True)
        if sts_images and st.button("🔍 Распознать данные", type="primary"):
            with st.spinner("Изучаю документы..."):
                try:
                    images_pil = [Image.open(img) for img in sts_images]
                    prompt = f"Это фото техпаспорта КР. Верни СТРОГО в формате JSON без markdown ключи: customer, region, district, aymak, street_address, car_model, reg_num, vin, tech_passport, year, color, engine_vol, body_type. Справочник регионов: {json.dumps(KG_REGIONS, ensure_ascii=False)}"
                    response = gemini_client.models.generate_content(
                        model='gemini-2.5-flash',
                        contents=[prompt] + images_pil
                    )
                    raw_json = re.sub(r'^```json\s*|\s*```$', '', response.text.strip(), flags=re.MULTILINE).strip()
                    data = json.loads(raw_json)
                    
                    if data.get("region") in KG_REGIONS:
                        st.session_state["region_select"] = data["region"]
                        district_val = data.get("district", "")
                        st.session_state["district_select"] = district_val if district_val in KG_REGIONS[data["region"]] else KG_REGIONS[data["region"]][0]
                    for key in ["customer", "aymak", "street_address", "car_model", "reg_num", "vin", "tech_passport", "year", "color", "engine_vol", "body_type"]:
                        raw_val = data.get(key, "")
                        str_val = "" if raw_val is None else str(raw_val)
                        if key == "aymak": st.session_state["aymak_input"] = str_val
                        else: st.session_state[key] = str_val
                    st.success("✅ Данные распознаны!")
                    st.rerun() 
                except Exception as e:
                    st.error(f"❌ Ошибка распознавания: {e}")

col_hdr1, col_hdr2 = st.columns([4, 1])
with col_hdr1: st.header("1. Ввод данных")
with col_hdr2: st.button("🧹 Очистить форму", on_click=clear_fields, width='stretch')

df_preview = get_cached_preview()
df_db = get_cached_db()

col1, col2 = st.columns(2)
with col1:
    report_num = st.text_input("Номер отчета:", key="report_num")
    contract_num = st.text_input("Номер договора:", key="contract_num")
    date_ocenki = st.text_input("Дата оценки:", key="date_ocenki") 
    date_otcheta = st.text_input("Дата отчета (для реестра):", key="date_otcheta", value=datetime.now().strftime("%d.%m.%Y"))
    customer = st.text_input("ФИО Заказчика:", key="customer")
    c_geo1, c_geo2, c_geo3 = st.columns(3)
    with c_geo1: selected_region = st.selectbox("Область / Город:", list(KG_REGIONS.keys()), key="region_select")
    with c_geo2: selected_district = st.selectbox("Район / Округ:", KG_REGIONS[selected_region], key="district_select")
    with c_geo3: aymak = st.text_input("Село:", key="aymak_input")
    street_detail = st.text_input("Улица, дом, квартира:", key="street_address")
    full_address = f"Кыргызская Республика, {selected_region}, {selected_district}, {aymak.strip()}, {street_detail.strip()}".replace(" ,", ",").strip(", ")
    
    sum_num = st.text_input("Сумма ущерба цифрами (для титульного листа):", key="sum_num")
    gen_sum = ""
    if sum_num:
        try: gen_sum = num2words(int(float(re.sub(r'[^\d.]', '', sum_num.replace(",", ".")))), lang='ru').lower()
        except: pass
    sum_words = st.text_input("Сумма ущерба прописью:", value=gen_sum)

with col2:
    car_model = st.text_input("Марка, модель:", key="car_model")
    reg_num = st.text_input("Гос. номер:", key="reg_num")
    vin = st.text_input("VIN код:", key="vin")
    tech_passport = st.text_input("Тех. паспорт №:", key="tech_passport")
    year = st.text_input("Год выпуска:", key="year")
    engine_vol = st.text_input("Объем ДВС:", key="engine_vol")
    color = st.text_input("Цвет кузова:", key="color")
    c_in1, c_in2 = st.columns(2)
    with c_in1: body_type = st.text_input("Тип кузова:", key="body_type")
    with c_in2: steering = st.selectbox("Руль:", ["Левый руль", "Правый руль"], key="steering")
    service_cost = st.text_input("💰 Стоимость услуги (заработок):", key="service_cost")

dup_warnings = find_duplicate_matches(
    dfs=[("Живой отчет для шефа", df_preview), ("База_проверок", df_db)],
    checks=[
        ("Номер отчета", report_num),
        ("ФИО Заказчика", customer),
        ("Тех. паспорт", tech_passport),
        ("Гос. номер", reg_num),
        ("VIN код", vin),
    ]
)
if dup_warnings:
    st.warning("🔁 Похоже, такой отчет уже создавался ранее — проверь, чтобы случайно не задвоить:\n\n" + "\n\n".join(dup_warnings))

st.header("2. Описание повреждений и ремонта")

# --- ВЕРНУЛ ВСЕ ШАБЛОНЫ НА МЕСТО ---
DAMAGE_TEMPLATES = {
    "--- Выберите шаблон ---": "",
    "[Кузов] Передняя часть": "При осмотре установлены повреждения передней части кузова: деформация бампера, повреждение облицовочных элементов, смещение/деформация навесных деталей, нарушение ЛКП.",
    "[Кузов] Задняя часть": "Выявлены повреждения задней части кузова: деформация бампера, повреждение крышки багажника/фонарей, нарушение геометрии сопряжений, повреждение ЛКП.",
    "[Кузов] Боковая часть": "Установлены повреждения боковой части кузова: деформация дверей/крыльев, повреждение навесных элементов, нарушение ЛКП.",
    "[Кузов] Силовые элементы": "Имеются признаки деформации силовых элементов кузова (лонжерон/панель), требующие восстановительных работ с последующим контролем геометрии.",
    "[Оптика] Фара (трещина/разрушение)": "Блок-фара передняя (указать сторону): сквозное разрушение (трещина) рассеивателя.",
    "[Оптика] Фара (царапины)": "Блок-фара передняя (указать сторону): глубокие царапины и потертости рассеивателя.",
    "[Оптика] Фара (крепления)": "Блок-фара передняя (указать сторону): излом элементов крепления корпуса.",
    "[Стекла] Лобовое (трещина)": "Стекло ветровое: линейная трещина в зоне видимости водителя (или: в зоне работы стеклоочистителей).",
    "[Стекла] Лобовое (скол)": "Стекло ветровое: скол типа «звезда» (или «бычий глаз») с развивающимися трещинами.",
    "[Стекла] Боковое (царапины)": "Стекло передней/задней двери (указать сторону): царапины (задиры) на внешней поверхности.",
    "[Стекла] Боковое (разрушение)": "Стекло передней/задней двери (указать сторону): разрушение элемента (отсутствует).",
    "[Стекла] Заднее (седан)": "Стекло задка: разрушение элемента / глубокие царапины.",
    "[Стекла] Заднее (хэтчбек/внедорожник)": "Стекло двери задка (крышки багажника): повреждение нитей обогрева / разрушение."
}

REPAIR_TEMPLATES = {
    "--- Выберите шаблон ---": "",
    "[Кузов] Стандартные работы": "Для восстановления требуется выполнить комплекс слесарно-кузовных, рихтовочных и малярно-окрасочных работ с применением расходных материалов, с последующей сборкой и регулировкой навесных элементов.",
    "[Оптика] Замена фары": "Демонтаж, монтаж (замена) блок-фары передней (указать сторону) в сборе.",
    "[Стекла] Лобовое стекло (база)": "Замена стекла ветрового (вклейка) с использованием комплекта однокомпонентного полиуретанового клея.",
    "[Стекла] Лобовое стекло (+датчики)": "Замена стекла ветрового (вклейка) с использованием комплекта однокомпонентного полиуретанового клея и переустановкой датчика дождя/камеры слежения.",
    "[Стекла] Боковое стекло": "Снятие обивки двери, очистка внутренней полости от осколков, замена стекла двери.",
    "[Стекла] Заднее стекло": "Замена стекла задка (вклейка) с подключением элементов обогрева."
}

if "damage_text" not in st.session_state: st.session_state.damage_text = DEFAULT_DAMAGE_SUFFIX
if "repair_text" not in st.session_state: st.session_state.repair_text = f"{REPAIR_TEMPLATES['[Кузов] Стандартные работы']}\n{DEFAULT_REPAIR_SUFFIX}"

def add_to_damage():
    selected = st.session_state.dmg_selector
    if selected and DAMAGE_TEMPLATES[selected]:
        st.session_state.damage_text = (st.session_state.damage_text.replace(DEFAULT_DAMAGE_SUFFIX, DAMAGE_TEMPLATES[selected] + "\n" + DEFAULT_DAMAGE_SUFFIX) if DEFAULT_DAMAGE_SUFFIX in st.session_state.damage_text else st.session_state.damage_text + "\n" + DAMAGE_TEMPLATES[selected])

def add_to_repair():
    selected = st.session_state.rep_selector
    if selected and REPAIR_TEMPLATES[selected]:
        st.session_state.repair_text = (st.session_state.repair_text.replace(DEFAULT_REPAIR_SUFFIX, REPAIR_TEMPLATES[selected] + "\n" + DEFAULT_REPAIR_SUFFIX) if DEFAULT_REPAIR_SUFFIX in st.session_state.repair_text else st.session_state.repair_text + "\n" + REPAIR_TEMPLATES[selected])

col_dmg, col_rep = st.columns(2)
with col_dmg:
    st.selectbox("Конструктор осмотра:", list(DAMAGE_TEMPLATES.keys()), key="dmg_selector")
    st.button("➕ Добавить в осмотр", on_click=add_to_damage, width='stretch')
with col_rep:
    st.selectbox("Конструктор ремонта:", list(REPAIR_TEMPLATES.keys()), key="rep_selector")
    st.button("➕ Добавить в ремонт", on_click=add_to_repair, width='stretch')

damage_desc = st.text_area("Характеристика повреждений:", key="damage_text", height=150)
repair_desc = st.text_area("Требуемый ремонт:", key="repair_text", height=150)

def format_text(text):
    rt = RichText()
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    for i, line in enumerate(lines): rt.add(line + '\n') if i < len(lines) - 1 else rt.add(line)
    return rt

st.header("3. Смета и дополнительные контакты")
st.info("💡 Загрузите файл со сметой (.docx) от шефа. Приложение само распознает таблицы и автоматически посчитает ИТОГО для пункта Согласования.")
calc_report_doc = st.file_uploader("Загрузите файл с расчетами (Смета .docx)", type="docx")

st.markdown("**Точечные добавки (если нужно дополнить фирменные примечания под таблицами):**")
c_ex1, c_ex2, c_ex3 = st.columns(3)
with c_ex1: extra_services = st.text_area("Доп. контакты для Услуг:", key="extra_services", height=100)
with c_ex2: extra_parts = st.text_area("Доп. ссылки для Запчастей:", key="extra_parts", height=100)
with c_ex3: extra_materials = st.text_area("Доп. данные для Материалов:", key="extra_materials", height=100)

st.header("4. Приложение: Фотоотчет")
photo_report_doc = st.file_uploader("Загрузите готовый Фотоотчет (.docx)", type="docx")

if template_source is not None:
    if st.button("СГЕНЕРИРОВАТЬ ИТОГОВЫЙ ОТЧЕТ", type="primary", width='stretch'):
        try:
            doc = DocxTemplate(template_source if isinstance(template_source, str) else template_source.seek(0) or template_source)
            
            subdoc_photo = doc.new_subdoc(photo_report_doc) if photo_report_doc else "Фотоотчет не приложен."
            
            approval_text = ""
            if calc_report_doc:
                smart_buffer, approval_text = process_smart_calc_tables(calc_report_doc, extra_services, extra_parts, extra_materials)
                subdoc_calc = doc.new_subdoc(smart_buffer) if smart_buffer else "Ошибка обработки таблиц."
            else:
                subdoc_calc = "Файл со сметой не был приложен."

            context = {
                "REPORT_NUM": report_num, "CONTRACT_NUM": contract_num, "DATE": date_ocenki,
                "CUSTOMER_NAME": customer, "ADDRESS": full_address, "CAR_MODEL": car_model,
                "REG_NUM": reg_num, "VIN": vin, "TECH_PASSPORT": tech_passport,
                "YEAR": year, "ENGINE_VOL": engine_vol, "COLOR": color, "BODY_TYPE": body_type,
                "STEERING": steering, "TOTAL_SUM_NUM": sum_num, "TOTAL_SUM_WORDS": sum_words,
                "DAMAGE_DESC": format_text(damage_desc), "REPAIR_DESC": format_text(repair_desc),
                "CALC_TABLES": subdoc_calc,
                "APPROVAL_BLOCK": format_text(approval_text),
                "PHOTO_TABLE": subdoc_photo 
            }
            doc.render(context)
            force_font_everywhere(doc.docx)
            
            buffer = io.BytesIO()
            doc.save(buffer)
            file_bytes = buffer.getvalue()
            file_name = f"{reg_num.strip() or 'Без_номера'}.docx"
            
            with st.spinner("Отправка в Telegram..."):
                tg_status = upload_to_telegram(file_bytes, file_name)
            
            if append_to_google_sheets([report_num, car_model, reg_num, date_ocenki, date_otcheta, service_cost, tg_status], [report_num, reg_num, vin, tech_passport, date_otcheta, tg_status]):
                get_cached_preview.clear(); get_cached_db.clear()
                st.success("✅ Отчет создан, таблицы обработаны, суммы подсчитаны и файл отправлен в Telegram!")
            
            st.download_button(f"📥 СКАЧАТЬ ОТЧЕТ ({file_name})", file_bytes, file_name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", width='stretch')
        except Exception as e:
            st.error(f"Произошла ошибка: {e}")

st.sidebar.title("📊 Живой отчет для шефа")
if df_preview is not None and not df_preview.empty:
    # Рендерим таблицу вручную через HTML, минуя pyarrow —
    # на сервере (pyarrow 25 + Python 3.14) st.dataframe/st.table
    # может падать с сегфолтом при смешанных типах данных из Google Sheets.
    try:
        html_table = df_preview.to_html(index=False, escape=True, na_rep="")
        st.sidebar.markdown(
            f'<div style="overflow-x:auto; font-size:13px;">{html_table}</div>',
            unsafe_allow_html=True
        )
    except Exception as e:
        st.sidebar.error(f"⚠️ Не удалось отобразить таблицу: {e}")
